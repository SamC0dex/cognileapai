from __future__ import annotations

import copy
import posixpath
import re
import shutil
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from lxml import etree


SAMPLE = Path(r"C:\Users\swami\Downloads\report ICVGTDP FINAL.docx")
SOURCE = Path(r"C:\Users\swami\Downloads\CogniLeapAI_Final_Project_Report_V3.docx")
OUT = Path(r"C:\Users\swami\Downloads\CogniLeapAI_Final_Project_Report_APPROVED_FORMAT.docx")

OLD_TITLE = "Automated Image Captioning and Voice Generation"
NEW_TITLE = (
    "CogniLeapAI: An Adaptive Personalized Learning Platform with AI-Powered "
    "Content Generation, Active Recall, and Intelligent Study Planning"
)

REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
OFFICE_REL_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
PKG_REL = f"{{{REL_NS}}}Relationship"
R_ID = f"{{{OFFICE_REL_NS}}}id"
R_EMBED = f"{{{OFFICE_REL_NS}}}embed"
R_LINK = f"{{{OFFICE_REL_NS}}}link"


def clear_body_keep_section(doc: Document) -> None:
    body = doc._element.body
    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)


def copy_body(source: Document, dest: Document) -> None:
    dest_body = dest._element.body
    dest_sect = dest_body.sectPr
    for child in source._element.body:
        if child.tag == qn("w:sectPr"):
            continue
        cloned = copy.deepcopy(child)
        for sect in cloned.xpath(".//w:sectPr"):
            parent = sect.getparent()
            if parent is not None:
                parent.remove(sect)
        dest_body.insert(dest_body.index(dest_sect), cloned)


def replace_in_paragraph(paragraph, old: str, new: str) -> None:
    if old not in paragraph.text:
        return
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)


def normalize_text(paragraph) -> None:
    text = paragraph.text.strip()
    upper = text.upper()

    if not text:
        return

    # Preserve table-specific style and alignment where the source already uses tables.
    if paragraph._p.getparent().tag.endswith("tc}"):
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            run.font.size = Pt(11)
        return

    is_chapter = upper.startswith("CHAPTER")
    is_front_heading = upper in {
        "CERTIFICATE",
        "PROJECT APPROVAL SHEET",
        "ACKNOWLEDGEMENT",
        "ABSTRACT",
        "TABLE OF CONTENTS",
        "LIST OF ABBREVIATIONS",
        "LIST OF FIGURES",
        "LIST OF TABLES",
        "REFERENCES",
    } or upper.startswith("ANNEXURE")

    if is_chapter or is_front_heading:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.bold = True
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            run.font.size = Pt(14 if is_chapter else 16)
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(12)
        return

    # Numbered section headings, e.g. 1.1 Detailed Problem Definition.
    parts = text.split(" ", 1)
    if parts and parts[0].replace(".", "").isdigit() and "." in parts[0]:
        for run in paragraph.runs:
            run.bold = True
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            run.font.size = Pt(12)
        paragraph.paragraph_format.space_before = Pt(10)
        paragraph.paragraph_format.space_after = Pt(6)
        return

    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        if run.font.size is None or run.font.size.pt < 10 or run.font.size.pt > 13:
            run.font.size = Pt(12)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(6)


def set_paragraph_text(paragraph, text: str) -> None:
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def normalize_heading_text(paragraph) -> None:
    text = paragraph.text.strip()
    if not text:
        return
    upper = text.upper()
    replacements = {
        "INDEX": "TABLE OF CONTENTS",
        "ABSTRACT": "ABSTRACT",
        "ACKNOWLEDGEMENT": "ACKNOWLEDGEMENT",
        "ACKNOWLEDGMENT": "ACKNOWLEDGEMENT",
        "REFERENCES": "REFERENCES",
        "LIST OF FIGURES": "LIST OF FIGURES",
        "LIST OF TABLES": "LIST OF TABLES",
        "LIST OF ABBREVIATIONS": "LIST OF ABBREVIATIONS",
        "CERTIFICATE": "CERTIFICATE",
        "PROJECT APPROVAL SHEET": "PROJECT APPROVAL SHEET",
    }
    if upper in replacements and text != replacements[upper]:
        set_paragraph_text(paragraph, replacements[upper])
        return
    match = re.match(r"^\s*CHAPTER\s+(\d+)\s*:\s*(.+?)\s*$", text, flags=re.I)
    if match:
        set_paragraph_text(paragraph, f"CHAPTER {int(match.group(1)):02d}\n{match.group(2).strip().upper()}")


def force_chapter_page_breaks(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip().upper()
        if re.match(r"^CHAPTER\s+0?[2-8]\b", text):
            paragraph.paragraph_format.page_break_before = True


def add_body_first_line_indents(doc: Document) -> None:
    in_body = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        upper = text.upper()
        if upper.startswith("CHAPTER 01"):
            in_body = True
        if not in_body or not text:
            continue
        if paragraph._p.getparent().tag.endswith("tc}"):
            continue
        if upper.startswith("CHAPTER") or upper.startswith("ANNEXURE") or upper == "REFERENCES":
            continue
        if re.match(r"^\d+(\.\d+)+\s+", text):
            continue
        if text.startswith(("Table ", "Figure ")):
            continue
        if paragraph.style.name in {"List Paragraph", "figure caption"}:
            continue
        paragraph.paragraph_format.first_line_indent = Inches(0.5)


def normalize_tables(doc: Document) -> None:
    for table in doc.tables:
        table.autofit = True
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip().upper() == "INDEX":
                        set_paragraph_text(paragraph, "TABLE OF CONTENTS")
                    for run in paragraph.runs:
                        run.font.name = "Times New Roman"
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                        run.font.size = Pt(10)


def replace_headers_footers(doc: Document) -> None:
    replacements = {
        OLD_TITLE: "CogniLeapAI",
        "Automated image captioning and voice generation": "CogniLeapAI",
        "Automated Image Captioning and Voice Generation": "CogniLeapAI",
    }
    for section in doc.sections:
        for container in (section.header, section.footer):
            for paragraph in container.paragraphs:
                for old, new in replacements.items():
                    replace_in_paragraph(paragraph, old, new)
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def apply_sample_page_setup(doc: Document, sample: Document) -> None:
    src = sample.sections[0]
    for section in doc.sections:
        section.page_height = src.page_height
        section.page_width = src.page_width
        section.top_margin = src.top_margin
        section.bottom_margin = src.bottom_margin
        section.left_margin = src.left_margin
        section.right_margin = src.right_margin
        section.header_distance = src.header_distance
        section.footer_distance = src.footer_distance


def main() -> None:
    sample = Document(SAMPLE)
    dest = Document(SOURCE)
    apply_sample_page_setup(dest, sample)
    replace_headers_footers(dest)
    for paragraph in dest.paragraphs:
        replace_in_paragraph(paragraph, OLD_TITLE, NEW_TITLE)
        normalize_heading_text(paragraph)
        normalize_text(paragraph)
    normalize_tables(dest)
    force_chapter_page_breaks(dest)
    add_body_first_line_indents(dest)
    dest.save(OUT)
    replace_ooxml_text(OUT, {"INDEX": "TABLE OF CONTENTS"})
    print(OUT)


def _read_xml(zf: zipfile.ZipFile, name: str) -> etree._Element:
    return etree.fromstring(zf.read(name))


def _next_rid(existing: set[str]) -> str:
    i = 1
    while f"rId{i}" in existing:
        i += 1
    rid = f"rId{i}"
    existing.add(rid)
    return rid


def _unique_part_name(existing: set[str], original: str) -> str:
    if original not in existing:
        existing.add(original)
        return original
    folder, name = posixpath.split(original)
    stem, dot, ext = name.rpartition(".")
    if not dot:
        stem, ext = name, ""
    for i in range(1, 10000):
        candidate_name = f"{stem}_cognileap_{i}.{ext}" if ext else f"{stem}_cognileap_{i}"
        candidate = posixpath.join(folder, candidate_name)
        if candidate not in existing:
            existing.add(candidate)
            return candidate
    raise RuntimeError(f"Could not create unique part name for {original}")


def transplant_document_relationships(source_path: Path, dest_path: Path) -> None:
    with zipfile.ZipFile(source_path) as src_zip, zipfile.ZipFile(dest_path) as dst_zip:
        document_xml = _read_xml(dst_zip, "word/document.xml")
        source_rels = _read_xml(src_zip, "word/_rels/document.xml.rels")
        dest_rels = _read_xml(dst_zip, "word/_rels/document.xml.rels")
        ns = {"r": OFFICE_REL_NS}
        used_ids = set(document_xml.xpath("//@r:id | //@r:embed | //@r:link", namespaces=ns))
        existing_rel_ids = {rel.get("Id") for rel in dest_rels.findall(PKG_REL)}
        existing_names = set(dst_zip.namelist())
        id_map: dict[str, str] = {}
        extra_files: dict[str, bytes] = {}

        for rel in source_rels.findall(PKG_REL):
            old_id = rel.get("Id")
            if old_id not in used_ids:
                continue
            rel_type = rel.get("Type", "")
            if not any(kind in rel_type for kind in ("image", "chart", "diagram", "hyperlink")):
                continue
            new_id = _next_rid(existing_rel_ids)
            id_map[old_id] = new_id
            new_rel = etree.Element(PKG_REL)
            new_rel.set("Id", new_id)
            new_rel.set("Type", rel_type)
            if rel.get("TargetMode"):
                new_rel.set("TargetMode", rel.get("TargetMode"))
                new_rel.set("Target", rel.get("Target"))
            else:
                source_target = posixpath.normpath(posixpath.join("word", rel.get("Target")))
                dest_target = _unique_part_name(existing_names, source_target)
                extra_files[dest_target] = src_zip.read(source_target)
                new_rel.set("Target", posixpath.relpath(dest_target, "word"))
            dest_rels.append(new_rel)

        for old_id, new_id in id_map.items():
            for attr in (R_ID, R_EMBED, R_LINK):
                for elem in document_xml.iter():
                    if elem.get(attr) == old_id:
                        elem.set(attr, new_id)

        tmp = dest_path.with_suffix(".tmp.docx")
        with zipfile.ZipFile(dest_path) as in_zip, zipfile.ZipFile(
            tmp, "w", zipfile.ZIP_DEFLATED
        ) as out_zip:
            for item in in_zip.infolist():
                if item.filename == "word/document.xml":
                    out_zip.writestr(item, etree.tostring(document_xml, xml_declaration=True, encoding="UTF-8", standalone="yes"))
                elif item.filename == "word/_rels/document.xml.rels":
                    out_zip.writestr(item, etree.tostring(dest_rels, xml_declaration=True, encoding="UTF-8", standalone="yes"))
                else:
                    out_zip.writestr(item, in_zip.read(item.filename))
            for name, data in extra_files.items():
                out_zip.writestr(name, data)
        shutil.move(tmp, dest_path)


def replace_ooxml_text(docx_path: Path, replacements: dict[str, str]) -> None:
    tmp = docx_path.with_suffix(".textfix.docx")
    with zipfile.ZipFile(docx_path) as in_zip, zipfile.ZipFile(
        tmp, "w", zipfile.ZIP_DEFLATED
    ) as out_zip:
        for item in in_zip.infolist():
            data = in_zip.read(item.filename)
            if item.filename.startswith("word/") and item.filename.endswith(".xml"):
                try:
                    root = etree.fromstring(data)
                    changed = False
                    for node in root.xpath(".//*[local-name()='t']"):
                        if node.text in replacements:
                            node.text = replacements[node.text]
                            changed = True
                    if changed:
                        data = etree.tostring(
                            root,
                            xml_declaration=True,
                            encoding="UTF-8",
                            standalone="yes",
                        )
                except Exception:
                    pass
            out_zip.writestr(item, data)
    shutil.move(tmp, docx_path)


if __name__ == "__main__":
    main()
