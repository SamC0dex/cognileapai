from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml.ns import qn


DOCX = Path(r"C:\Users\swami\Downloads\CogniLeapAI_Final_Project_Report_APPROVED_FORMAT.docx")

SCREENSHOTS = [
    ("6.4", "Dashboard and Study Library Overview", r"C:\Users\swami\OneDrive\Pictures\Screenshots\Screenshot 2026-06-26 184821.png"),
    ("6.5", "Dashboard with File Panel and Generated Study Materials", r"C:\Users\swami\OneDrive\Pictures\Screenshots\Screenshot 2026-06-26 185555.png"),
    ("6.6", "Document-Aware Chat Interface with Selected PDF", r"C:\Users\swami\OneDrive\Pictures\Screenshots\Screenshot 2026-06-26 185926.png"),
    ("6.7", "Chat Model Selection and Multi-Model Routing Options", r"C:\Users\swami\OneDrive\Pictures\Screenshots\Screenshot 2026-06-26 191527.png"),
    ("6.8", "Study Guide Generated from Uploaded BI Document", r"C:\Users\swami\OneDrive\Pictures\Screenshots\Screenshot 2026-06-26 190426.png"),
    ("6.9", "Quiz Interface Generated from Study Material", r"C:\Users\swami\OneDrive\Pictures\Screenshots\Screenshot 2026-06-26 190043.png"),
    ("6.10", "Mind Map Visualization with Topic Detail Panel", r"C:\Users\swami\OneDrive\Pictures\Screenshots\Screenshot 2026-06-26 190256.png"),
    ("6.11", "Generated Flashcard Viewer", r"C:\Users\swami\OneDrive\Pictures\Screenshots\Screenshot 2026-06-26 190554.png"),
    ("6.12", "Active Recall Dashboard and Smart Review Planning", r"C:\Users\swami\OneDrive\Pictures\Screenshots\Screenshot 2026-06-26 185725.png"),
    ("6.13", "Active Recall Review Session with SM-2 Rating Controls", r"C:\Users\swami\OneDrive\Pictures\Screenshots\Screenshot 2026-06-26 185830.png"),
    ("6.14", "Encrypted API Key Management Screen", r"C:\Users\swami\OneDrive\Pictures\Screenshots\Screenshot 2026-06-26 190126.png"),
    ("6.15", "Default AI Provider and Model Settings", r"C:\Users\swami\OneDrive\Pictures\Screenshots\Screenshot 2026-06-26 190143.png"),
    ("6.16", "Dashboard Card View of Generated Learning Assets", r"C:\Users\swami\OneDrive\Pictures\Screenshots\Screenshot 2026-06-26 183139.png"),
]


def set_font(paragraph, size=12, bold=None, italic=None):
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(size)
        if bold is not None:
            run.bold = bold
        if italic is not None:
            run.italic = italic


def insert_paragraph_before(paragraph, text="", style=None):
    new_p = paragraph.insert_paragraph_before(text)
    if style:
        new_p.style = style
    return new_p


def main() -> None:
    doc = Document(DOCX)
    # Avoid duplicate insertion on rerun.
    for p in doc.paragraphs:
        if p.text.strip() == "6.6 Feature Screenshots":
            raise SystemExit("Feature screenshots section already exists; not inserting duplicates.")

    chapter7 = next(p for p in doc.paragraphs if p.text.strip().upper().startswith("CHAPTER 07"))

    h = insert_paragraph_before(chapter7, "6.6 Feature Screenshots")
    h.paragraph_format.page_break_before = True
    h.paragraph_format.first_line_indent = None
    h.paragraph_format.space_before = Pt(0)
    h.paragraph_format.space_after = Pt(8)
    set_font(h, 12, bold=True)

    intro = insert_paragraph_before(
        chapter7,
        "The following screenshots present the implemented CogniLeapAI feature set as tested during the final validation phase. They show the dashboard, document-aware chat, study tool generation, Active Recall review flow, and AI provider configuration screens used by the deployed system.",
    )
    intro.paragraph_format.first_line_indent = Inches(0.5)
    intro.paragraph_format.line_spacing = 1.5
    intro.paragraph_format.space_after = Pt(8)
    set_font(intro, 12)

    for fig_no, caption, path in SCREENSHOTS:
        img_path = Path(path)
        if not img_path.exists():
            raise FileNotFoundError(img_path)
        spacer = insert_paragraph_before(chapter7, "")
        spacer.paragraph_format.page_break_before = True
        image_p = insert_paragraph_before(chapter7, "")
        image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = image_p.add_run()
        # Landscape browser screenshots fit the text width; portrait screenshots stay narrower.
        width = Inches(6.35) if img_path.stat().st_size else Inches(6.35)
        if "190043" in img_path.name:
            width = Inches(3.7)
        run.add_picture(str(img_path), width=width)

        cap = insert_paragraph_before(chapter7, f"Figure {fig_no}: {caption}")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.first_line_indent = None
        cap.paragraph_format.space_before = Pt(4)
        cap.paragraph_format.space_after = Pt(8)
        set_font(cap, 12, bold=True, italic=True)

    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    main()
