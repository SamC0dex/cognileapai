from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


DOCX = Path(r"C:\Users\swami\Downloads\CogniLeapAI FInal Version.docx")
OUT = Path(r"C:\Users\swami\Downloads\CogniLeapAI FInal Version_FRONTMATTER_FIXED.docx")

TOC_LINES = [
    "List Of Abbreviation                                                                    vii",
    "List Of Figures                                                                           ix",
    "List Of Tables                                                                             x",
    "",
    "",
    "Chapter No.\t                                    Title                                                 Page No.",
    "",
    "",
    "01\t\tIntroduction\t1",
    "1.1\tDetailed Problem Definition\t1",
    "1.2\tJustification of Problem\t3",
    "1.3\tNeed for the New System\t4",
    "1.4\tAdvances Over Previous Systems\t5",
    "1.5\tPresently Available Systems\t6",
    "1.6\tPurpose of the System\t7",
    "1.7\tOrganization of the Report\t8",
    "02\t\tAnalysis and Literature Survey\t15",
    "2.1\tProject Plan\t15",
    "2.2\tRequirement Analysis\t17",
    "2.3\tTeam Structure\t20",
    "2.4\tLiterature Survey\t21",
    "03\t\tFramework and Architecture\t23",
    "3.1\tSoftware Requirement Specification (SRS)\t23",
    "3.2\tSystem Architecture\t27",
    "3.3\tRisk Assessment\t29",
    "3.4\tMajor Milestones\t30",
    "3.5\tDevelopment Paradigm\t31",
    "04\t\tModeling\t31",
    "4.1\tSystem Architecture and Functional Block Diagram\t31",
    "4.2\tUser Flow Diagram\t33",
    "4.3\tData Flow Diagrams\t34",
    "4.4\tEntity Relationship Diagram and Normalization\t37",
    "4.5\tUse Case Diagram\t39",
    "4.6\tClass Diagram\t40",
    "4.7\tSequence Diagram\t41",
    "4.8\tState Chart Diagram\t42",
    "4.9\tActivity Diagram\t43",
    "4.10\tComponent Diagram\t44",
    "4.11\tDeployment Diagram\t45",
    "4.12\tObject Diagram\t46",
    "05\t\tSoftware Requirements and Coding\t47",
    "5.1\tSoftware and Hardware Specifications\t47",
    "5.2\tProgramming Languages and Platform\t49",
    "5.3\tComponents and Tools\t50",
    "5.4\tAlgorithms and Flowcharts\t52",
    "5.5\tCoding Style and Standards\t56",
    "06\t\tTest Data Sets, Result and Analysis\t59",
    "6.1\tTest Plan\t59",
    "6.2\tTest Cases\t60",
    "6.3\tTest Results\t64",
    "6.4\tPerformance Test Results\t66",
    "6.5\tFormat Technical Reviews\t68",
    "6.6\tFeature Screenshots\t69",
    "07\t\tConfiguration Management Plan\t83",
    "7.1\tVersion Control\t83",
    "7.2\tConfiguration Items\t84",
    "7.3\tEnvironment Management\t85",
    "7.4\tBuild and Deployment\t86",
    "7.5\tSoftware Quality Assurance Plan\t87",
    "7.6\tWeekly Development Report\t88",
    "08\t\tConclusion and Future Scope\t90",
    "8.1\tConclusion\t90",
    "8.2\tFuture Scope\t91",
    "References\t93",
]

ABBREVIATIONS = [
    ("AI", "Artificial Intelligence"),
    ("API", "Application Programming Interface"),
    ("BYOK", "Bring Your Own Key"),
    ("CRUD", "Create, Read, Update, Delete"),
    ("DFD", "Data Flow Diagram"),
    ("ERD", "Entity Relationship Diagram"),
    ("FR", "Functional Requirement"),
    ("JSON", "JavaScript Object Notation"),
    ("LLM", "Large Language Model"),
    ("LMS", "Learning Management System"),
    ("MVP", "Minimum Viable Product"),
    ("NFR", "Non-Functional Requirement"),
    ("OAuth", "Open Authorization"),
    ("PDF", "Portable Document Format"),
    ("PKCE", "Proof Key for Code Exchange"),
    ("RLS", "Row Level Security"),
    ("SM-2", "SuperMemo 2 Spaced Repetition Algorithm"),
    ("SRS", "Software Requirement Specification"),
    ("SSE", "Server-Sent Events"),
    ("UI", "User Interface"),
    ("UML", "Unified Modeling Language"),
    ("URL", "Uniform Resource Locator"),
]

FULL_ABBREVIATIONS = [
    ("AI", "Artificial Intelligence"),
    ("API", "Application Programming Interface"),
    ("BYOK", "Bring Your Own Key"),
    ("CRUD", "Create, Read, Update, Delete"),
    ("DFD", "Data Flow Diagram"),
    ("EF", "Easiness Factor"),
    ("ER", "Entity Relationship"),
    ("FR", "Functional Requirement"),
    ("HTML", "HyperText Markup Language"),
    ("HTTP", "HyperText Transfer Protocol"),
    ("JSON", "JavaScript Object Notation"),
    ("LLM", "Large Language Model"),
    ("LMS", "Learning Management System"),
    ("MVP", "Minimum Viable Product"),
    ("NFR", "Non-Functional Requirement"),
    ("OAuth", "Open Authorization"),
    ("PDF", "Portable Document Format"),
    ("PKCE", "Proof Key for Code Exchange"),
    ("RLS", "Row Level Security"),
    ("SRS", "Software Requirement Specification"),
    ("SM-2", "SuperMemo 2 Spaced Repetition Algorithm"),
    ("SSE", "Server-Sent Events"),
    ("TDD", "Test-Driven Development"),
    ("TOC", "Table of Contents"),
    ("UI", "User Interface"),
    ("UML", "Unified Modeling Language"),
    ("URL", "Uniform Resource Locator"),
    ("UX", "User Experience"),
]


def set_paragraph_text(p, text: str) -> None:
    for run in p.runs:
        run.text = ""
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)


def set_run_font(p, size=None, bold=None):
    for run in p.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        if size:
            run.font.size = Pt(size)
        if bold is not None:
            run.bold = bold


def main() -> None:
    doc = Document(DOCX)

    # Replace the visible sample-style TOC paragraph block.
    start, end = 120, 181
    capacity = end - start + 1
    if len(TOC_LINES) > capacity:
        anchor = doc.paragraphs[end + 1]
        for _ in range(len(TOC_LINES) - capacity):
            p = anchor.insert_paragraph_before("")
            p.style = doc.paragraphs[end].style
    for offset, text in enumerate(TOC_LINES):
        p = doc.paragraphs[start + offset]
        set_paragraph_text(p, text)
        set_run_font(p, 12)
    for i in range(start + len(TOC_LINES), end + 1):
        set_paragraph_text(doc.paragraphs[i], "")

    # Replace only the pasted sample abbreviation row tables while preserving format.
    # In this document those are exactly tables 4 through 25. Table 26 and later are
    # real report content tables and must not be touched.
    for idx, table in enumerate(doc.tables[4:26], start=0):
        if idx >= len(ABBREVIATIONS):
            for row in table.rows:
                for cell in row.cells:
                    cell.text = ""
            continue
        abbr, full = ABBREVIATIONS[idx]
        if len(table.rows) and len(table.columns) >= 2:
            table.cell(0, 0).text = abbr
            table.cell(0, 1).text = full
            for cell in (table.cell(0, 0), table.cell(0, 1)):
                for p in cell.paragraphs:
                    set_run_font(p, 12)

    # Also keep the full structured abbreviation table current if Word chooses
    # to paginate/display it instead of the single-row pasted sample tables.
    if len(doc.tables) > 26 and len(doc.tables[26].columns) == 2:
        table = doc.tables[26]
        while len(table.rows) > 1:
            table._tbl.remove(table.rows[-1]._tr)
        table.cell(0, 0).text = "Abbreviation"
        table.cell(0, 1).text = "Full Form"
        for abbr, full in FULL_ABBREVIATIONS:
            cells = table.add_row().cells
            cells[0].text = abbr
            cells[1].text = full
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    set_run_font(p, 12)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
