from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


SRC = Path(r"C:\Users\swami\Downloads\CogniLeapAI FInal Version.docx")
OUT = Path(r"C:\Users\swami\Downloads\CogniLeapAI FInal Version_TOC_FIXED.docx")

TOC_ROWS = [
    ("", "List Of Abbreviation", "vii"),
    ("", "List Of Figures", "ix"),
    ("", "List Of Tables", "x"),
    ("", "", ""),
    ("Chapter No.", "Title", "Page No."),
    ("", "", ""),
    ("01", "Introduction", "1"),
    ("1.1", "Detailed Problem Definition", "1"),
    ("1.2", "Justification of Problem", "3"),
    ("1.3", "Need for the New System", "4"),
    ("1.4", "Advances Over Previous Systems", "5"),
    ("1.5", "Presently Available Systems", "6"),
    ("1.6", "Purpose of the System", "7"),
    ("1.7", "Organization of the Report", "8"),
    ("02", "Analysis and Literature Survey", "15"),
    ("2.1", "Project Plan", "15"),
    ("2.2", "Requirement Analysis", "17"),
    ("2.3", "Team Structure", "20"),
    ("2.4", "Literature Survey", "21"),
    ("03", "Framework and Architecture", "23"),
    ("3.1", "Software Requirement Specification (SRS)", "23"),
    ("3.2", "System Architecture", "27"),
    ("3.3", "Risk Assessment", "29"),
    ("3.4", "Major Milestones", "30"),
    ("3.5", "Development Paradigm", "31"),
    ("04", "Modeling", "31"),
    ("4.1", "System Architecture and Functional Block Diagram", "31"),
    ("4.2", "User Flow Diagram", "33"),
    ("4.3", "Data Flow Diagrams", "34"),
    ("4.4", "Entity Relationship Diagram and Normalization", "37"),
    ("4.5", "Use Case Diagram", "39"),
    ("4.6", "Class Diagram", "40"),
    ("4.7", "Sequence Diagram", "41"),
    ("4.8", "State Chart Diagram", "42"),
    ("4.9", "Activity Diagram", "43"),
    ("4.10", "Component Diagram", "44"),
    ("4.11", "Deployment Diagram", "45"),
    ("4.12", "Object Diagram", "46"),
    ("05", "Software Requirements and Coding", "47"),
    ("5.1", "Software and Hardware Specifications", "47"),
    ("5.2", "Programming Languages and Platform", "49"),
    ("5.3", "Components and Tools", "50"),
    ("5.4", "Algorithms and Flowcharts", "52"),
    ("5.5", "Coding Style and Standards", "56"),
    ("06", "Test Data Sets, Result and Analysis", "59"),
    ("6.1", "Test Plan", "59"),
    ("6.2", "Test Cases", "60"),
    ("6.3", "Test Results", "64"),
    ("6.4", "Performance Test Results", "66"),
    ("6.5", "Format Technical Reviews", "68"),
    ("6.6", "Feature Screenshots", "69"),
    ("07", "Configuration Management Plan", "83"),
    ("7.1", "Version Control", "83"),
    ("7.2", "Configuration Items", "84"),
    ("7.3", "Environment Management", "85"),
    ("7.4", "Build and Deployment", "86"),
    ("7.5", "Software Quality Assurance Plan", "87"),
    ("7.6", "Weekly Development Report", "88"),
    ("08", "Conclusion and Future Scope", "90"),
    ("8.1", "Conclusion", "90"),
    ("8.2", "Future Scope", "91"),
    ("", "References", "93"),
    ("", "Annexure A: Mathematical Model, Database Description, and Additional Information", "96"),
    ("", "Annexure B: Plagiarism Summary Report", "100"),
    ("", "Annexure C: Publication Details and Published Full-Length Papers", "102"),
]

FULL_ABBREVIATIONS = [
    ("AI", "Artificial Intelligence"),
    ("API", "Application Programming Interface"),
    ("BYOK", "Bring Your Own Key"),
    ("CRUD", "Create, Read, Update, Delete"),
    ("DFD", "Data Flow Diagram"),
    ("EF", "Easiness Factor"),
    ("ER", "Entity Relationship"),
    ("FR", "Functional Requirement"),
    ("HTML", "HyperText Markup Language"),
    ("HTTP", "HyperText Transfer Protocol"),
    ("JSON", "JavaScript Object Notation"),
    ("LLM", "Large Language Model"),
    ("LMS", "Learning Management System"),
    ("MVP", "Minimum Viable Product"),
    ("NFR", "Non-Functional Requirement"),
    ("OAuth", "Open Authorization"),
    ("PDF", "Portable Document Format"),
    ("PKCE", "Proof Key for Code Exchange"),
    ("RLS", "Row Level Security"),
    ("SM-2", "SuperMemo 2 Spaced Repetition Algorithm"),
    ("SRS", "Software Requirement Specification"),
    ("SSE", "Server-Sent Events"),
    ("UI", "User Interface"),
    ("UML", "Unified Modeling Language"),
    ("URL", "Uniform Resource Locator"),
]


def font_cell(cell, bold=False):
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.name = "Times New Roman"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            r.font.size = Pt(12)
            r.bold = bold


def set_cell(cell, text, bold=False):
    cell.text = text
    font_cell(cell, bold)


def rewrite_toc_table(doc: Document):
    table = doc.tables[1]
    while len(table.rows) < len(TOC_ROWS):
        table.add_row()
    while len(table.rows) > len(TOC_ROWS):
        table._tbl.remove(table.rows[-1]._tr)
    for row, values in zip(table.rows, TOC_ROWS):
        is_header = values == ("Chapter No.", "Title", "Page No.")
        is_major = bool(values[0].isdigit() and len(values[0]) == 2) or values[1].startswith(("References", "Annexure"))
        for i, value in enumerate(values):
            set_cell(row.cells[i], value, bold=(is_header or is_major))


def rewrite_abbrev_table(doc: Document):
    if len(doc.tables) <= 26:
        return
    table = doc.tables[26]
    if len(table.columns) != 2:
        return
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)
    set_cell(table.cell(0, 0), "Abbreviation", bold=True)
    set_cell(table.cell(0, 1), "Full Form", bold=True)
    for abbr, full in FULL_ABBREVIATIONS:
        cells = table.add_row().cells
        set_cell(cells[0], abbr)
        set_cell(cells[1], full)


def main():
    doc = Document(SRC)
    rewrite_toc_table(doc)
    rewrite_abbrev_table(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
