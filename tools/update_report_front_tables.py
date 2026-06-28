from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn


DOCX = Path(r"C:\Users\swami\Downloads\CogniLeapAI_Final_Project_Report_APPROVED_FORMAT.docx")
FRONT_OFFSET = 13

FIGURE_ROWS = [
    ("3.1", "High-Level System Architecture", "23"),
    ("4.1", "System Architecture and Component Layout", "31"),
    ("4.2", "Functional Block Diagram", "32"),
    ("4.3", "User Flow Diagram", "33"),
    ("4.4", "Data Flow Diagram - Level 0 (Context Diagram)", "34"),
    ("4.5", "Data Flow Diagram - Level 1", "35"),
    ("4.6", "Data Flow Diagram - Level 2", "36"),
    ("4.7", "Entity Relationship Diagram", "38"),
    ("4.8", "Use Case Diagram", "39"),
    ("4.9", "Class Diagram", "40"),
    ("4.10", "Sequence Diagram - Chat Streaming Pipeline", "41"),
    ("4.11", "State Chart Diagram - Active Recall Card Lifecycle", "42"),
    ("4.12", "Activity Diagram - Study Tools Generation", "43"),
    ("4.13", "Component Diagram", "44"),
    ("4.14", "Deployment Diagram", "45"),
    ("4.15", "Object Diagram", "46"),
    ("5.1", "SM-2 Spaced Repetition Algorithm Flowchart", "53"),
    ("5.2", "Multi-Provider AI Routing Flowchart", "54"),
    ("5.3", "Priority-Scored Card Scheduling Flowchart", "55"),
    ("6.1", "Cumulative Test Pass Rate by Phase", "66"),
    ("6.2", "Retention Curve Comparison: With and Without SM-2", "67"),
    ("6.3", "Study Efficiency: Manual vs CogniLeapAI", "67"),
    ("6.4", "Dashboard and Study Library Overview", "70"),
    ("6.5", "Dashboard with File Panel and Generated Study Materials", "71"),
    ("6.6", "Document-Aware Chat Interface with Selected PDF", "72"),
    ("6.7", "Chat Model Selection and Multi-Model Routing Options", "73"),
    ("6.8", "Study Guide Generated from Uploaded BI Document", "74"),
    ("6.9", "Quiz Interface Generated from Study Material", "75"),
    ("6.10", "Mind Map Visualization with Topic Detail Panel", "76"),
    ("6.11", "Generated Flashcard Viewer", "77"),
    ("6.12", "Active Recall Dashboard and Smart Review Planning", "78"),
    ("6.13", "Active Recall Review Session with SM-2 Rating Controls", "79"),
    ("6.14", "Encrypted API Key Management Screen", "80"),
    ("6.15", "Default AI Provider and Model Settings", "81"),
    ("6.16", "Dashboard Card View of Generated Learning Assets", "82"),
]

TOC_ROWS = [
    ("", "List of Figures", ""),
    ("", "List of Tables", ""),
    ("", "List of Abbreviations", ""),
    ("1", "Chapter 1: Introduction", "1"),
    ("1.1", "Detailed Problem Definition", "1"),
    ("1.2", "Justification of Problem", "3"),
    ("1.3", "Need for the New System", "4"),
    ("1.4", "Advances Over Previous Systems", "5"),
    ("1.5", "Presently Available Systems", "6"),
    ("1.6", "Purpose of the System", "7"),
    ("1.7", "Organization of the Report", "8"),
    ("2", "Chapter 2: Analysis and Literature Survey", "15"),
    ("2.1", "Project Plan", "15"),
    ("2.2", "Requirement Analysis", "17"),
    ("2.3", "Team Structure", "20"),
    ("2.4", "Literature Survey", "21"),
    ("3", "Chapter 3: Framework and Architecture", "23"),
    ("3.1", "Software Requirement Specification (SRS)", "23"),
    ("3.2", "System Architecture", "27"),
    ("3.3", "Risk Assessment", "29"),
    ("3.4", "Major Milestones", "30"),
    ("3.5", "Development Paradigm", "31"),
    ("4", "Chapter 4: Modeling", "31"),
    ("4.1", "System Architecture and Functional Block Diagram", "31"),
    ("4.2", "User Flow Diagram", "33"),
    ("4.3", "Data Flow Diagrams", "34"),
    ("4.4", "Entity Relationship Diagram and Normalization", "37"),
    ("4.5", "Use Case Diagram", "39"),
    ("4.6", "Class Diagram", "40"),
    ("4.7", "Sequence Diagram", "41"),
    ("4.8", "State Chart Diagram", "42"),
    ("4.9", "Activity Diagram", "43"),
    ("4.10", "Component Diagram", "44"),
    ("4.11", "Deployment Diagram", "45"),
    ("4.12", "Object Diagram", "46"),
    ("5", "Chapter 5: Software Requirements and Coding", "47"),
    ("5.1", "Software and Hardware Specifications", "47"),
    ("5.2", "Programming Languages and Platform", "49"),
    ("5.3", "Components and Tools", "50"),
    ("5.4", "Algorithms and Flowcharts", "52"),
    ("5.5", "Coding Style and Standards", "56"),
    ("6", "Chapter 6: Test Data Sets, Result and Analysis", "59"),
    ("6.1", "Test Plan", "59"),
    ("6.2", "Test Cases", "60"),
    ("6.3", "Test Results", "64"),
    ("6.4", "Performance Test Results", "66"),
    ("6.5", "Format Technical Reviews", "68"),
    ("6.6", "Feature Screenshots", "69"),
    ("7", "Chapter 7: Configuration Management Plan", "83"),
    ("7.1", "Version Control", "83"),
    ("7.2", "Configuration Items", "84"),
    ("7.3", "Environment Management", "85"),
    ("7.4", "Build and Deployment", "86"),
    ("7.5", "Software Quality Assurance Plan", "87"),
    ("7.6", "Weekly Development Report", "88"),
    ("8", "Chapter 8: Conclusion and Future Scope", "90"),
    ("8.1", "Conclusion", "90"),
    ("8.2", "Future Scope", "91"),
    ("", "References", "93"),
    ("", "Annexure A: Mathematical Model, Database Description, and Additional Information", "96"),
    ("", "Annexure B: Plagiarism Summary Report", "100"),
    ("", "Annexure C: Publication Details and Published Full-Length Papers", "102"),
]


def set_cell(cell, text, bold=False):
    cell.text = text
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            run.font.size = Pt(11)
            run.bold = bold


def rewrite_table(table, rows, header):
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)
    for i, text in enumerate(header):
        set_cell(table.rows[0].cells[i], text, bold=True)
    for row in rows:
        cells = table.add_row().cells
        is_major = row[0].isdigit() or row[1].startswith(("References", "Annexure"))
        for i, value in enumerate(row):
            set_cell(cells[i], value, bold=is_major)


def main():
    doc = Document(DOCX)
    rewrite_table(doc.tables[0], TOC_ROWS, ("Sr. No.", "Chapter / Title", "Page No."))
    rewrite_table(doc.tables[1], FIGURE_ROWS, ("Figure No.", "Description", "Page No."))
    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    main()
