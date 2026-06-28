from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


SRC = Path(r"C:\Users\swami\Downloads\CogniLeapAI FInal Version.docx")
OUT = Path(r"C:\Users\swami\Downloads\CogniLeapAI FInal Version_SAMPLE_TOC_FIXED.docx")

TOC_LINES = [
    "List Of Abbreviation                                                                    vii",
    "List Of Figures                                                                           ix",
    "List Of Tables                                                                             x",
    "",
    "",
    "Chapter No.\t                                    Title                                                 Page No.",
    "",
    "",
    "01\t\tIntroduction\t1",
    "1.1\tDetailed Problem Definition\t1",
    "1.2\tJustification of Problem\t3",
    "1.3\tNeed for the New System\t4",
    "1.4\tAdvances Over Previous Systems\t5",
    "1.5\tPresently Available Systems\t6",
    "1.6\tPurpose of the System\t7",
    "1.7\tOrganization of the Report\t8",
    "02\t\tAnalysis and Literature Survey\t15",
    "2.1\tProject Plan\t15",
    "2.2\tRequirement Analysis\t17",
    "2.3\tTeam Structure\t20",
    "2.4\tLiterature Survey\t21",
    "03\t\tFramework and Architecture\t23",
    "3.1\tSoftware Requirement Specification (SRS)\t23",
    "3.2\tSystem Architecture\t27",
    "3.3\tRisk Assessment\t29",
    "3.4\tMajor Milestones\t30",
    "3.5\tDevelopment Paradigm\t31",
    "04\t\tModeling\t31",
    "4.1\tSystem Architecture and Functional Block Diagram\t31",
    "4.2\tUser Flow Diagram\t33",
    "4.3\tData Flow Diagrams\t34",
    "4.4\tEntity Relationship Diagram and Normalization\t37",
    "4.5\tUse Case Diagram\t39",
    "4.6\tClass Diagram\t40",
    "4.7\tSequence Diagram\t41",
    "4.8\tState Chart Diagram\t42",
    "4.9\tActivity Diagram\t43",
    "4.10\tComponent Diagram\t44",
    "4.11\tDeployment Diagram\t45",
    "4.12\tObject Diagram\t46",
    "05\t\tSoftware Requirements and Coding\t47",
    "5.1\tSoftware and Hardware Specifications\t47",
    "5.2\tProgramming Languages and Platform\t49",
    "5.3\tComponents and Tools\t50",
    "5.4\tAlgorithms and Flowcharts\t52",
    "5.5\tCoding Style and Standards\t56",
    "06\t\tTest Data Sets, Result and Analysis\t59",
    "6.1\tTest Plan\t59",
    "6.2\tTest Cases\t60",
    "6.3\tTest Results\t64",
    "6.4\tPerformance Test Results\t66",
    "6.5\tFormat Technical Reviews\t68",
    "6.6\tFeature Screenshots\t69",
    "07\t\tConfiguration Management Plan\t83",
    "7.1\tVersion Control\t83",
    "7.2\tConfiguration Items\t84",
    "7.3\tEnvironment Management\t85",
    "7.4\tBuild and Deployment\t86",
    "7.5\tSoftware Quality Assurance Plan\t87",
    "7.6\tWeekly Development Report\t88",
    "08\t\tConclusion and Future Scope\t90",
    "8.1\tConclusion\t90",
    "8.2\tFuture Scope\t91",
    "References\t93",
    "Annexure A: Mathematical Model, Database Description, and Additional Information\t96",
    "Annexure B: Plagiarism Summary Report\t100",
    "Annexure C: Publication Details and Published Full-Length Papers\t102",
]

ABBREVIATIONS = [
    ("AI", "Artificial Intelligence"),
    ("API", "Application Programming Interface"),
    ("BYOK", "Bring Your Own Key"),
    ("CRUD", "Create, Read, Update, Delete"),
    ("DFD", "Data Flow Diagram"),
    ("EF", "Easiness Factor"),
    ("ER", "Entity Relationship"),
    ("FR", "Functional Requirement"),
    ("HTML", "HyperText Markup Language"),
    ("HTTP", "HyperText Transfer Protocol"),
    ("JSON", "JavaScript Object Notation"),
    ("LLM", "Large Language Model"),
    ("LMS", "Learning Management System"),
    ("MVP", "Minimum Viable Product"),
    ("NFR", "Non-Functional Requirement"),
    ("OAuth", "Open Authorization"),
    ("PDF", "Portable Document Format"),
    ("PKCE", "Proof Key for Code Exchange"),
    ("RLS", "Row Level Security"),
    ("SM-2", "SuperMemo 2 Spaced Repetition Algorithm"),
    ("SRS", "Software Requirement Specification"),
    ("SSE", "Server-Sent Events"),
    ("UI", "User Interface"),
    ("UML", "Unified Modeling Language"),
    ("URL", "Uniform Resource Locator"),
]


def set_text(p, text):
    for r in p.runs:
        r.text = ""
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)
    for r in p.runs:
        r.font.name = "Times New Roman"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        r.font.size = Pt(12)


def font_cell(cell, bold=False):
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.name = "Times New Roman"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            r.font.size = Pt(12)
            r.bold = bold


def main():
    doc = Document(SRC)

    start, original_end = 120, 181
    capacity = original_end - start + 1
    if len(TOC_LINES) > capacity:
        anchor = doc.paragraphs[original_end + 1]
        for _ in range(len(TOC_LINES) - capacity):
            p = anchor.insert_paragraph_before("")
            p.style = doc.paragraphs[original_end].style
    for i, line in enumerate(TOC_LINES):
        set_text(doc.paragraphs[start + i], line)
    for i in range(start + len(TOC_LINES), original_end + 1):
        set_text(doc.paragraphs[i], "")

    # Remove the duplicate no-border table TOC so only the sample paragraph TOC remains.
    if len(doc.tables) > 1:
        doc.tables[1]._element.getparent().remove(doc.tables[1]._element)

    # Refresh the full abbreviation table only. After deleting table 1, this is table 25.
    for table in doc.tables:
        if len(table.columns) == 2 and table.cell(0, 0).text.strip() == "Abbreviation":
            while len(table.rows) > 1:
                table._tbl.remove(table.rows[-1]._tr)
            table.cell(0, 0).text = "Abbreviation"
            table.cell(0, 1).text = "Full Form"
            font_cell(table.cell(0, 0), True)
            font_cell(table.cell(0, 1), True)
            for abbr, full in ABBREVIATIONS:
                cells = table.add_row().cells
                cells[0].text = abbr
                cells[1].text = full
                font_cell(cells[0])
                font_cell(cells[1])
            break

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
